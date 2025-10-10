import os
import pandas as pd
import json
import hashlib
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document
from docx import Document as DocxDocument

load_dotenv()

# --- 1. DEFINE PATHS ---
DATA_PATH = "data/slp_knowledge_base/"
DB_PATH = "data/slp_vector_db/"
CLINICAL_DATA_PATH = "data/processed_clinical_data/"

# Additional PDF folder paths
PDF_FOLDERS = {
    'clinical_data': 'clinical_data',
    'case_books': 'case_books', 
    'peer_reviewed_articles': 'peer_reviewed_articles_simucase',
    'sample_iep': 'sample_iep',
    'school_policy': 'schoolpolicy_and_guidance'
}

class ComprehensiveSLPProcessor:
    def __init__(self):
        self.grouping_strategies = {
            1: "Speech sound disorder",
            2: "Articulation disorders", 
            3: "Phonological disorders",
            4: "Language disorders",
            5: "Receptive language disorders",
            6: "Expressive language disorders", 
            7: "Pragmatics",
            8: "Fluency",
            9: "Childhood apraxia of speech"
        }
        
        self.grouping_combinations = {
            "speech_sound_focus": [1, 2, 3, 9],
            "language_focus": [4, 5, 6, 7],
            "fluency_compatible": [4, 5, 6, 7, 8, 9],
            "mixed_speech_language": [1, 2, 3, 4, 5, 6]
        }
    
    def load_pdf_collection(self, folder_name, folder_path):
        """Load PDFs from a specific folder with appropriate metadata."""
        full_path = os.path.join(DATA_PATH, folder_path)
        documents = []
        
        if not os.path.exists(full_path):
            print(f"   ⚠ Folder not found: {full_path}")
            return documents
        
        try:
            pdf_loader = DirectoryLoader(
                full_path,
                glob="**/*.pdf",
                loader_cls=PyPDFLoader,
                use_multithreading=True
            )
            pdf_docs = pdf_loader.load()
            
            # Add enhanced metadata based on folder type
            for doc in pdf_docs:
                # Extract filename for better source identification
                source_file = os.path.basename(doc.metadata.get('source', ''))
                
                # Enhance content based on folder type
                enhanced_content = self._enhance_content_by_type(doc.page_content, folder_name, source_file)
                
                enhanced_doc = Document(
                    page_content=enhanced_content,
                    metadata={
                        'source': source_file,
                        'source_type': folder_name,
                        'full_path': doc.metadata.get('source', ''),
                        'page': doc.metadata.get('page', 0),
                        'collection': folder_name
                    }
                )
                documents.append(enhanced_doc)
            
            print(f"   ✓ Loaded {len(pdf_docs)} PDFs from {folder_name}")
            return documents
            
        except Exception as e:
            print(f"   ❌ Error loading PDFs from {folder_name}: {e}")
            return documents
    
    def _enhance_content_by_type(self, content, folder_type, filename):
        """Add contextual information based on document type."""
        
        context_prefix = ""
        
        if folder_type == 'case_books':
            context_prefix = f"""
            CASE STUDY REFERENCE - {filename}
            This is a clinical case study that can inform SLP practice and grouping decisions.
            Relevant for understanding intervention strategies and student outcomes.
            
            Content: """
            
        elif folder_type == 'peer_reviewed_articles':
            context_prefix = f"""
            PEER-REVIEWED RESEARCH - {filename}
            This is evidence-based research that supports clinical decision-making in SLP practice.
            Use this information to validate intervention approaches and grouping strategies.
            
            Content: """
            
        elif folder_type == 'sample_iep':
            context_prefix = f"""
            IEP SAMPLE/TEMPLATE - {filename}
            This document provides examples of IEP goals and documentation standards.
            Relevant for goal writing and progress monitoring in SLP services.
            Based on grouping strategies, students with similar goals can be grouped together.
            
            Content: """
            
        elif folder_type == 'school_policy':
            context_prefix = f"""
            SCHOOL POLICY/GUIDANCE - {filename}
            This document contains institutional policies and procedures for SLP services.
            Important for compliance and service delivery standards.
            
            Content: """
            
        elif folder_type == 'clinical_data':
            context_prefix = f"""
            CLINICAL REFERENCE - {filename}
            This document contains clinical information relevant to SLP practice.
            
            Content: """
        
        return context_prefix + content
    
    def extract_clinical_data_from_docx(self, file_path):
        """Extract structured data from SLP data sheets (same as before)."""
        try:
            doc = DocxDocument(file_path)
            documents = []
            
            student_info = {}
            goal_info = {}
            session_notes = []
            
            # Process tables to extract structured data
            for table_idx, table in enumerate(doc.tables):
                # Extract student info from first table
                if table_idx == 0:
                    for row_idx, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        
                        if len(cells) >= 4 and any('student' in cell.lower() for cell in cells):
                            if row_idx + 1 < len(table.rows):
                                data_row = [cell.text.strip() for cell in table.rows[row_idx + 1].cells]
                                if len(data_row) >= 4:
                                    name_hash = hashlib.md5(data_row[0].encode()).hexdigest()[:6]
                                    student_info = {
                                        'anonymous_name': f"Student_{name_hash}",
                                        'room': data_row[1],
                                        'teacher': f"Teacher_{hashlib.md5(data_row[2].encode()).hexdigest()[:6]}",
                                        'grade': data_row[3]
                                    }
                
                # Extract goal and session data (same logic as before)
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    for cell in cells:
                        if cell.startswith('From ') and '/' in cell:
                            goal_info['date'] = cell
                        elif len(cell) > 50 and ('ARD' in cell or 'annual' in cell.lower()):
                            anonymized_goal = cell.replace('Beckem', '[STUDENT]')
                            goal_info['text'] = anonymized_goal
                
                # Extract session notes
                header_found = False
                date_col = time_col = data_col = None
                
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    
                    if not header_found and len(cells) >= 4:
                        headers_lower = [cell.lower() for cell in cells]
                        if 'date' in headers_lower and 'data' in headers_lower:
                            header_found = True
                            date_col = headers_lower.index('date')
                            data_col = headers_lower.index('data')
                            if 'time' in headers_lower:
                                time_col = headers_lower.index('time')
                            continue
                    
                    if header_found and len(cells) > max(date_col, data_col):
                        date_text = cells[date_col] if date_col < len(cells) else ""
                        time_text = cells[time_col] if time_col is not None and time_col < len(cells) else ""
                        data_text = cells[data_col] if data_col < len(cells) else ""
                        
                        if date_text and data_text and len(data_text) > 3:
                            full_date = f"{date_text} {time_text}".strip()
                            session_notes.append({
                                'date': full_date,
                                'note': data_text
                            })
            
            # Create documents (same logic as before but with enhanced context)
            file_hash = hashlib.md5(file_path.encode()).hexdigest()[:8]
            
            if student_info:
                grade = student_info.get('grade', '4')
                profile_content = f"""
                CLINICAL DATA SHEET - Student Profile
                Anonymous ID: {student_info['anonymous_name']}
                Grade: {grade}
                Room: {student_info.get('room', '')}
                Teacher: {student_info.get('teacher', '')}
                
                GROUPING ANALYSIS:
                This Grade {grade} student can be grouped with students in grades {max(int(grade)-2, 0)} through {int(grade)+2}.
                Based on the pragmatic communication goal, this student has Type 7 (Pragmatics) disorder.
                Compatible grouping combinations: Types 4, 5, 6, 7 (language focus) or with Types 8, 9.
                
                CLINICAL CONTEXT:
                Student receives SLP services for pragmatic communication difficulties.
                Primary intervention focus: Topic maintenance using AAC (Augmentative Alternative Communication).
                This case demonstrates real therapy outcomes and progress tracking.
                """
                
                documents.append(Document(
                    page_content=profile_content,
                    metadata={
                        'source': f"clinical_data_sheet_{file_hash}",
                        'type': 'student_profile',
                        'grade': grade,
                        'disorder_type': 'pragmatics',
                        'collection': 'clinical_data_sheets'
                    }
                ))
            
            if goal_info:
                goal_content = f"""
                CLINICAL DATA SHEET - Annual IEP Goal
                Goal Date: {goal_info.get('date', 'From 10/17')}
                Goal Text: {goal_info.get('text', '')}
                
                CLINICAL ANALYSIS:
                This goal targets pragmatic communication skills (Type 7 disorder).
                Intervention approach: AAC usage for topic maintenance.
                Success criteria: 3 conversational turns, 4/5 trials, across 2 data collection sessions.
                
                GROUPING IMPLICATIONS:
                Students with similar pragmatic goals can be grouped together.
                Compatible with language disorder students (Types 4, 5, 6, 7).
                Can also be grouped with fluency (Type 8) or childhood apraxia (Type 9) students.
                """
                
                documents.append(Document(
                    page_content=goal_content,
                    metadata={
                        'source': f"clinical_data_sheet_{file_hash}",
                        'type': 'annual_goal',
                        'disorder_type': 'pragmatics',
                        'intervention_type': 'AAC_topic_maintenance',
                        'collection': 'clinical_data_sheets'
                    }
                ))
            
            for i, note in enumerate(session_notes):
                if note['note'].strip() and note['note'].strip().lower() not in ['absent', 'absent, sick']:
                    note_content = f"""
                    CLINICAL DATA SHEET - Therapy Session Note
                    Date: {note['date']}
                    Session Data: {note['note']}
                    
                    INTERVENTION ANALYSIS:
                    This session demonstrates AAC usage and pragmatic skill development.
                    Student Grade: {student_info.get('grade', '4')}
                    Focus: Topic maintenance, conversational turns, AAC device navigation.
                    
                    CLINICAL SIGNIFICANCE:
                    Real therapy data showing progress on pragmatic communication goals.
                    Demonstrates effective intervention strategies for similar students.
                    Relevant for grouping students with comparable needs and goals.
                    """
                    
                    documents.append(Document(
                        page_content=note_content,
                        metadata={
                            'source': f"clinical_data_sheet_{file_hash}",
                            'type': 'session_note',
                            'date': note['date'],
                            'session_number': i + 1,
                            'collection': 'clinical_data_sheets'
                        }
                    ))
            
            return documents
            
        except Exception as e:
            print(f"Error processing clinical data sheet {file_path}: {e}")
            return []

# --- 2. LOAD ALL DOCUMENTS ---
print("="*70)
print("LOADING COMPREHENSIVE SLP KNOWLEDGE BASE")
print("="*70)

processor = ComprehensiveSLPProcessor()
all_documents = []

# Load PDFs from all specified folders
print("\n📚 Loading PDF collections...")
for folder_name, folder_path in PDF_FOLDERS.items():
    print(f"\n{folder_name.upper().replace('_', ' ')}:")
    pdf_docs = processor.load_pdf_collection(folder_name, folder_path)
    all_documents.extend(pdf_docs)

# Load clinical data sheets (DOCX format)
print(f"\n📋 CLINICAL DATA SHEETS:")
clinical_documents = []
clinical_data_path = os.path.join(DATA_PATH, "clinical_data_sheets")

if os.path.exists(clinical_data_path):
    docx_files = [f for f in os.listdir(clinical_data_path) if f.endswith('.docx') and not f.startswith('~$')]
    print(f"   Found {len(docx_files)} clinical data sheet files")
    
    for file_name in docx_files:
        full_path = os.path.join(clinical_data_path, file_name)
        clinical_docs = processor.extract_clinical_data_from_docx(full_path)
        clinical_documents.extend(clinical_docs)
        print(f"   ✓ Processed {file_name}: {len(clinical_docs)} documents extracted")
else:
    print(f"   ⚠ Clinical data sheets directory not found: {clinical_data_path}")

all_documents.extend(clinical_documents)

# Add comprehensive grouping strategies document
print(f"\n📖 ADDING SLP GROUPING STRATEGIES...")
grouping_doc = Document(
    page_content="""
    SLP GROUPING STRATEGIES - COMPREHENSIVE GUIDELINES
    
    GROUP COMPOSITION RULES:
    • Group Size: 2-4 students (2-3 most common, 4 students rare)
    • Grade Levels: Pre-K through 12th grade
    • Grade Compatibility: Maximum 2 grade levels difference between group members
    
    GRADE LEVEL COMPATIBILITY:
    • Pre-K: Compatible with K, 1st grade
    • Kindergarten: Compatible with Pre-K, 1st, 2nd grade  
    • Each grade level: Compatible with ±2 grade levels
    • Example: 4th grade can group with 2nd, 3rd, 4th, 5th, 6th grades
    
    COMMUNICATION DISORDER TYPES:
    1. Speech Sound Disorder: Articulation and phonology disorders
    2. Articulation Disorders: Specific residual errors (/r/, /l/, /th/)
    3. Phonological Disorders: Sound patterns (deletion, fronting, stopping)
    4. Language Disorders: Receptive and expressive language disorders
    5. Receptive Language Disorders: Understanding difficulties
    6. Expressive Language Disorders: Expression difficulties (semantics, morphology, syntax)
    7. Pragmatics: Social communication, topic maintenance, turn-taking
    8. Fluency: Stuttering or cluttering
    9. Childhood Apraxia of Speech: Motor planning difficulties
    
    OPTIMAL GROUPING COMBINATIONS:
    • Speech Sound Focus: Types 1, 2, 3, 9 together
    • Language Focus: Types 4, 5, 6, 7 together  
    • Fluency Integration: Type 8 with 4, 5, 6, 7, 9
    • Mixed Groups: Types 1, 2, 3 with 4, 5, 6
    
    CLINICAL EXAMPLE (from data sheet):
    Grade 4 student with pragmatics (Type 7) working on topic maintenance:
    - Can group with grades 2-6
    - Compatible with language disorders (Types 4, 5, 6, 7)
    - Can include fluency (Type 8) or apraxia (Type 9) students
    - AAC intervention strategies work well in mixed groups
    """,
    metadata={
        'source': 'slp_grouping_strategies_comprehensive',
        'type': 'clinical_guidelines',
        'collection': 'grouping_protocols'
    }
)

all_documents.append(grouping_doc)

# Document loading summary
print(f"\n{'='*70}")
print("DOCUMENT LOADING SUMMARY")
print(f"{'='*70}")

collection_counts = {}
for doc in all_documents:
    collection = doc.metadata.get('collection', 'unknown')
    collection_counts[collection] = collection_counts.get(collection, 0) + 1

for collection, count in collection_counts.items():
    print(f"{collection.replace('_', ' ').title()}: {count} documents")

print(f"\nTOTAL DOCUMENTS: {len(all_documents)}")

if not all_documents:
    print(f"\n❌ No documents found. Please check your folder structure in {DATA_PATH}")
else:
    # --- 3. CHUNK THE DOCUMENTS ---
    print(f"\n{'='*70}")
    print("CHUNKING DOCUMENTS FOR OPTIMAL EMBEDDING")
    print(f"{'='*70}")
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,  # Slightly larger for comprehensive content
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len
    )
    
    texts = text_splitter.split_documents(all_documents)
    print(f"✓ Split {len(all_documents)} documents into {len(texts)} chunks")
    
    # Analyze chunk distribution
    chunk_collections = {}
    for chunk in texts:
        collection = chunk.metadata.get('collection', 'unknown')
        chunk_collections[collection] = chunk_collections.get(collection, 0) + 1
    
    print("\nChunk distribution by collection:")
    for collection, count in chunk_collections.items():
        print(f"   {collection.replace('_', ' ').title()}: {count} chunks")

    # --- 4. CREATE EMBEDDINGS ---
    print(f"\n{'='*70}")
    print("CREATING EMBEDDINGS WITH OPENAI")
    print(f"{'='*70}")
    print("This may take several minutes for a comprehensive knowledge base...")
    
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-small"
    )

    # --- 5. STORE IN VECTOR DATABASE ---
    print(f"\n{'='*70}")
    print("STORING IN CHROMADB VECTOR DATABASE")
    print(f"{'='*70}")
    
    vectorstore = Chroma.from_documents(
        documents=texts,
        embedding=embeddings,
        persist_directory=DB_PATH
    )
    
    # Test queries for different document types
    print(f"\n🔍 Testing database with sample queries...")
    
    test_queries = [
        "What are the grouping strategies for grade 4 students with pragmatic disorders?",
        "Show me IEP goal examples for topic maintenance",
        "What does research say about AAC interventions?",
        "What are school policies for SLP service delivery?"
    ]
    
    for query in test_queries:
        results = vectorstore.similarity_search(query, k=2)
        print(f"   Query: '{query[:50]}...'")
        print(f"   Retrieved {len(results)} relevant chunks from {set(r.metadata.get('collection', 'unknown') for r in results)}")
    
    # Final success message
    print(f"\n{'='*70}")
    print("✅ COMPREHENSIVE SLP VECTOR DATABASE COMPLETE!")
    print(f"{'='*70}")
    print(f"📊 Database Statistics:")
    print(f"   • Total documents processed: {len(all_documents)}")
    print(f"   • Total chunks created: {len(texts)}")
    print(f"   • Database location: {DB_PATH}")
    print(f"   • Collections included: {len(collection_counts)}")
    
    print(f"\n🎯 Your comprehensive database includes:")
    print(f"   • Clinical case studies and reference books")
    print(f"   • Peer-reviewed research articles") 
    print(f"   • Sample IEP goals and documentation")
    print(f"   • School policies and guidance documents")
    print(f"   • Real clinical session data and progress notes")
    print(f"   • Evidence-based grouping strategies")
    
    print(f"\n🔍 Ready for complex queries like:")
    print(f"   • 'What does research support for AAC interventions in pragmatics?'")
    print(f"   • 'Show me IEP examples similar to my Grade 4 topic maintenance goal'")
    print(f"   • 'What are school policy requirements for SLP documentation?'")
    print(f"   • 'Find case studies of students with similar profiles'")
    
    print(f"\n💡 Your RAG system can now provide:")
    print(f"   • Evidence-based recommendations from peer-reviewed research")
    print(f"   • Policy-compliant documentation guidance")
    print(f"   • Real clinical examples and case studies")
    print(f"   • Comprehensive grouping decision support")
    
    print(f"\n{'='*70}")